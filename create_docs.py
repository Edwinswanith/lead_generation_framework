from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_documentation():
    doc = Document()
    
    # Create styles
    styles = doc.styles
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Add title
    title = doc.add_heading('Lead Generation System - Your Smart Business Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Simple Introduction
    doc.add_heading('What is This System?', level=1)
    doc.add_paragraph(
        'Think of the Lead Generation System as your smart business assistant. It helps you find '
        'new potential customers and reaches out to them automatically. Instead of spending hours '
        'searching for business contacts and writing emails, this system does it all for you, '
        'saving you time and effort.'
    )
    
    # What It Does
    doc.add_heading('What Can It Do For You?', level=1)
    
    capabilities = [
        ('Find New Customers',
         'The system searches for companies that might be interested in your services. It looks '
         'at what they do, how big they are, and finds the right person to contact.'),
        
        ('Write Personal Emails',
         'Instead of sending the same email to everyone, the system creates unique messages for '
         'each company. It mentions their specific business and how you can help them.'),
        
        ('Keep Track of Everything',
         'You can see exactly what\'s happening at any time. The system shows you which companies '
         'it found, which emails were sent, and helps you follow up.')
    ]
    
    for title, description in capabilities:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)

    # How It Helps
    doc.add_heading('How Does It Make Your Life Easier?', level=1)
    
    benefits = [
        ('Saves Your Time',
         'What usually takes hours of work is done in minutes. You can focus on other important '
         'tasks while the system does the research and writing for you.'),
        
        ('Never Misses an Opportunity',
         'The system carefully checks each company to make sure they\'re a good fit for your '
         'business. It helps you find opportunities you might have missed.'),
        
        ('Keeps Everything Organized',
         'All your potential customer information, emails, and progress are kept in one place. '
         'No more scattered notes or lost contacts.')
    ]
    
    for title, description in benefits:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)

    # How to Use It
    doc.add_heading('How Do You Use It?', level=1)
    doc.add_paragraph(
        'Using the system is as easy as 1-2-3:'
    )
    
    steps = [
        ('Step 1: Set Up',
         'Just log in to the website and enter your email details. The system will keep these '
         'safe and use them to send emails on your behalf.'),
        
        ('Step 2: Start Finding Leads',
         'Click the "Generate Leads" button and tell the system what kind of companies you\'re '
         'looking for. Then watch as it finds potential customers for you.'),
        
        ('Step 3: Send Emails',
         'Review the companies the system found, and when you\'re ready, click "Send Emails". '
         'The system will create and send personalized emails to each company.')
    ]
    
    for title, description in steps:
        doc.add_heading(title, level=2)
        doc.add_paragraph(description)

    # Features Made Simple
    doc.add_heading('What Makes It Special?', level=1)
    
    features = [
        'Smart searching that finds the right companies for your business',
        'Automatic email writing that sounds personal and professional',
        'Easy-to-use website that shows you everything that\'s happening',
        'Safe storage of all your information and contacts',
        'Automatic progress updates so you know what\'s happening'
    ]
    
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature)

    # Real Benefits
    doc.add_heading('What Results Can You Expect?', level=1)
    doc.add_paragraph(
        'Here\'s what our system helps you achieve:'
    )
    
    results = [
        'Find more potential customers in less time',
        'Send professional, personalized emails without writing them yourself',
        'Keep track of all your leads in one place',
        'Save hours of work every week',
        'Reach more companies than you could on your own'
    ]
    
    for result in results:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(result)

    # Getting Help
    doc.add_heading('Need Help?', level=1)
    doc.add_paragraph(
        'We\'re here to help you succeed. If you ever have questions or need assistance, you can:'
    )
    
    help_options = [
        'Ask for help through the website',
        'Email our support team',
        'Check our simple guide that explains everything step by step',
        'Watch our helpful tutorial videos'
    ]
    
    for option in help_options:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(option)

    # Success Tips
    doc.add_heading('Tips for Success', level=1)
    doc.add_paragraph(
        'Here are some simple tips to get the most out of the system:'
    )
    
    tips = [
        'Start with a specific type of company you want to reach',
        'Check the leads the system finds before sending emails',
        'Use the system regularly to keep your pipeline full',
        'Keep track of which emails work best',
        'Follow up with companies that show interest'
    ]
    
    for tip in tips:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tip)

    # Save the document
    doc.save('Lead_Generation_Guide.docx')

if __name__ == '__main__':
    create_documentation()